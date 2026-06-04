from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Пути к файлам
vkr_path = "vkr/build/VKR_Villa_Jaconda.docx"

print("=" * 80)
print("ПОЛНЫЙ АНАЛИЗ ТРЕБОВАНИЙ ОФОРМЛЕНИЯ ВКР")
print("=" * 80)

doc_vkr = Document(vkr_path)

print(f"\n[СТРУКТУРА_ДОКУМЕНТА]")
print(f"Параграфов: {len(doc_vkr.paragraphs)}")
print(f"Таблиц: {len(doc_vkr.tables)}")
print(f"Разделов: {len(doc_vkr.sections)}")

# Анализ параметров страницы
print(f"\n[ПАРАМЕТРЫ_СТРАНИЦЫ]")
for i, section in enumerate(doc_vkr.sections):
    print(f"Раздел {i}:")
    print(f"  Размер страницы: {section.page_width.cm:.1f} см x {section.page_height.cm:.1f} см")
    print(f"  Верхнее поле: {section.top_margin.cm:.2f} см")
    print(f"  Нижнее поле: {section.bottom_margin.cm:.2f} см")
    print(f"  Левое поле: {section.left_margin.cm:.2f} см")
    print(f"  Правое поле: {section.right_margin.cm:.2f} см")

# Глубокий анализ форматирования
print(f"\n[АНАЛИЗ_ФОРМАТИРОВАНИЯ]")

# Собираем статистику по всему документу
font_sizes = {}
font_names = {}
styles_count = {}
bold_count = 0
italic_count = 0
underline_count = 0
line_spacing_values = {}
alignment_count = {}

for para in doc_vkr.paragraphs:
    # Подсчет стилей
    style_name = para.style.name if para.style else "No Style"
    styles_count[style_name] = styles_count.get(style_name, 0) + 1
    
    # Выравнивание
    alignment = para.alignment
    alignment_count[str(alignment)] = alignment_count.get(str(alignment), 0) + 1
    
    # Интервалы между строками
    if para.paragraph_format.line_spacing:
        ls = str(para.paragraph_format.line_spacing)
        line_spacing_values[ls] = line_spacing_values.get(ls, 0) + 1
    
    # Анализ текста в параграфе
    for run in para.runs:
        if run.text.strip():
            # Размер шрифта
            if run.font.size:
                pt_size = run.font.size.pt
                font_sizes[pt_size] = font_sizes.get(pt_size, 0) + 1
            
            # Имя шрифта
            if run.font.name:
                font_names[run.font.name] = font_names.get(run.font.name, 0) + 1
            
            # Жирный текст
            if run.font.bold:
                bold_count += 1
            
            # Курсив
            if run.font.italic:
                italic_count += 1
            
            # Подчеркивание
            if run.font.underline:
                underline_count += 1

print(f"\nИспользуемые стили параграфов:")
for style, count in sorted(styles_count.items(), key=lambda x: x[1], reverse=True):
    print(f"  {style}: {count}")

print(f"\nРазмеры шрифтов (встречаемость):")
for size, count in sorted(font_sizes.items(), key=lambda x: x[1], reverse=True):
    print(f"  {size}pt: {count} раз(а)")

print(f"\nИмена шрифтов (встречаемость):")
for name, count in sorted(font_names.items(), key=lambda x: x[1], reverse=True):
    if name:
        print(f"  {name}: {count} раз(а)")

print(f"\nФорматирование текста:")
print(f"  Жирный текст: {bold_count} элемент(ов)")
print(f"  Курсив: {italic_count} элемент(ов)")
print(f"  Подчеркивание: {underline_count} элемент(ов)")

print(f"\nИнтервалы между строками:")
for spacing, count in sorted(line_spacing_values.items()):
    print(f"  {spacing}: {count} параграфов")

# Анализ первого и последнего заголовков
print(f"\n[ЗАГОЛОВКИ]")
heading_count = 0
for para in doc_vkr.paragraphs:
    if para.style and 'Heading' in para.style.name:
        heading_count += 1

print(f"Всего заголовков (всех уровней): {heading_count}")

# Анализ таблиц
print(f"\n[ТАБЛИЦЫ]")
print(f"Количество таблиц: {len(doc_vkr.tables)}")
if len(doc_vkr.tables) > 0:
    for i, table in enumerate(doc_vkr.tables[:3]):
        print(f"  Таблица {i+1}: {len(table.rows)} строк, {len(table.columns)} столбцов")

# Выявленные проблемы
print(f"\n[ВЫЯВЛЕННЫЕ_ПРОБЛЕМЫ]")
problems = []

# Проверка полей
section = doc_vkr.sections[0]
if section.right_margin.cm < 1.4:
    problems.append(f"ОШИБКА: Правое поле {section.right_margin.cm:.2f} см (требуется минимум 1.5 см)")
if section.left_margin.cm < 2.9:
    problems.append(f"ПРЕДУПРЕЖДЕНИЕ: Левое поле {section.left_margin.cm:.2f} см (требуется 3.0 см)")

# Проверка размера шрифта
if 12 not in font_sizes and 14 not in font_sizes:
    problems.append(f"ПРЕДУПРЕЖДЕНИЕ: Основной размер шрифта не 12pt или 14pt. Обнаружены: {list(font_sizes.keys())}")
elif 14 in font_sizes and 12 not in font_sizes:
    problems.append(f"ИНФОРМАЦИЯ: Используется размер шрифта 14pt (вместо стандартного 12pt)")

# Проверка стилей
if "Normal (Web)" in styles_count:
    problems.append(f"ПРЕДУПРЕЖДЕНИЕ: Обнаружены стили 'Normal (Web)' ({styles_count['Normal (Web)']} шт.) - могут быть несовместимы со ГОСТ")

if not problems:
    problems.append("ОШИБОК НЕ ОБНАРУЖЕНО")

for i, problem in enumerate(problems, 1):
    print(f"  {i}. {problem}")

# ГОСТ требования (для справки)
print(f"\n[РЕКОМЕНДАЦИИ_ГОСТ]")
gost_reqs = [
    "Шрифт: Times New Roman или Arial",
    "Размер шрифта основного текста: 12 пункт",
    "Размер шрифта заголовков: не менее 12 пункт",
    "Поля: верхнее и нижнее - 2.0 см, левое - 3.0 см, правое - 1.5 см",
    "Интервал между строками: 1.5 или одинарный",
    "Выравнивание: по ширине",
    "Отступ первой строки: 1.25 см",
]
for i, req in enumerate(gost_reqs, 1):
    print(f"  {i}. {req}")
