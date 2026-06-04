from docx import Document
import os
from docx.shared import Pt, Cm, RGBColor

# Путь к файлам
template_path = "ВКР ПРИН редакция 2023_ФТИ.doc"
vkr_path = "vkr/build/VKR_Villa_Jaconda.docx"

print("=" * 80)
print("АНАЛИЗ ТРЕБОВАНИЙ ПО ОФОРМЛЕНИЮ")
print("=" * 80)

# Попытаемся прочитать шаблон
if os.path.exists(template_path):
    try:
        print(f"\n[TEMPLATE] Читаем шаблон: {template_path}")
        doc_template = Document(template_path)
        
        print(f"\n[OK] Шаблон успешно загружен")
        print(f"   Параграфов: {len(doc_template.paragraphs)}")
        print(f"   Таблиц: {len(doc_template.tables)}")
        
        # Выведем первые параграфы
        print(f"\n[TEMPLATE_CONTENT] Первые параграфы шаблона:")
        for i, para in enumerate(doc_template.paragraphs[:40]):
            if para.text.strip():
                style = para.style.name if para.style else "No Style"
                level = para.paragraph_format.outline_level if para.paragraph_format else 0
                print(f"   [{i}] ({style}): {para.text[:100]}")
                
        # Информация о странице
        print(f"\n[TEMPLATE_PAGE_INFO]:")
        for i, section in enumerate(doc_template.sections):
            print(f"   Раздел {i}:")
            print(f"     Размер: {section.page_width.cm:.1f} см x {section.page_height.cm:.1f} см")
            print(f"     Поля: верх={section.top_margin.cm:.1f}, низ={section.bottom_margin.cm:.1f}, "
                  f"лево={section.left_margin.cm:.1f}, право={section.right_margin.cm:.1f}")
                
    except Exception as e:
        print(f"[ERROR] Ошибка при чтении шаблона: {e}")
else:
    print(f"[ERROR] Файл шаблона не найден: {template_path}")

# Теперь читаем VKR файл
print(f"\n\n{'=' * 80}")
print(f"[VKR] Читаем VKR документ: {vkr_path}")
doc_vkr = Document(vkr_path)

print(f"\n[OK] VKR документ загружен")
print(f"   Параграфов: {len(doc_vkr.paragraphs)}")
print(f"   Таблиц: {len(doc_vkr.tables)}")
print(f"   Разделов: {len(doc_vkr.sections)}")

# Информация о странице VKR
print(f"\n[VKR_PAGE_INFO] Информация о странице:")
for i, section in enumerate(doc_vkr.sections):
    print(f"   Раздел {i}:")
    print(f"     Размер: {section.page_width.cm:.1f} см x {section.page_height.cm:.1f} см")
    print(f"     Поля: верх={section.top_margin.cm:.1f}, низ={section.bottom_margin.cm:.1f}, "
          f"лево={section.left_margin.cm:.1f}, право={section.right_margin.cm:.1f}")

# Анализ стилей в VKR
print(f"\n[VKR_STYLES] Используемые стили в VKR (первые 40 параграфов):")
styles_used = {}
for i, para in enumerate(doc_vkr.paragraphs[:40]):
    if para.text.strip():
        style = para.style.name if para.style else "No Style"
        if style not in styles_used:
            styles_used[style] = 0
        styles_used[style] += 1
        
        # Информация о шрифте
        if para.runs:
            run = para.runs[0]
            font_name = run.font.name if run.font.name else "Default"
            font_size = run.font.size if run.font.size else "Default"
            bold = "BOLD" if run.font.bold else ""
            print(f"   [{i}] Style: {style:20} | Font: {font_name:15} | Size: {font_size} {bold}")
            print(f"        Text: {para.text[:70]}")

print(f"\n[VKR_SUMMARY] Сводка стилей:")
for style, count in styles_used.items():
    print(f"   {style}: {count} параграфов")

# Проверка наличие нумерации и структуры
print(f"\n[VKR_STRUCTURE] Структура документа VKR:")
heading_count = 0
for para in doc_vkr.paragraphs:
    if para.style and 'Heading' in para.style.name:
        heading_count += 1
        print(f"   {para.style.name}: {para.text[:70]}")
        if heading_count >= 20:
            print(f"   ... (всего заголовков найдено {heading_count}+)")
            break

# Подробный анализ шрифта
print(f"\n[VKR_FONTS] Анализ шрифтов в первых параграфах VKR:")
for i, para in enumerate(doc_vkr.paragraphs[:30]):
    if para.text.strip():
        if para.runs:
            run = para.runs[0]
            font_info = []
            if run.font.name:
                font_info.append(f"шрифт: {run.font.name}")
            if run.font.size:
                pt_size = run.font.size.pt
                font_info.append(f"размер: {pt_size}pt")
            if run.font.bold:
                font_info.append("BOLD")
            print(f"   {para.style.name:15} | {', '.join(font_info):40} | {para.text[:50]}")

# Анализ интервалов между строками
print(f"\n[VKR_SPACING] Анализ интервалов (первые параграфы):")
for i, para in enumerate(doc_vkr.paragraphs[11:20]):
    pf = para.paragraph_format
    spacing_info = f"до: {pf.space_before.pt if pf.space_before else 0}pt, "
    spacing_info += f"после: {pf.space_after.pt if pf.space_after else 0}pt"
    if pf.line_spacing:
        spacing_info += f", междустрок: {pf.line_spacing}"
    print(f"   {spacing_info} | {para.text[:50]}")
