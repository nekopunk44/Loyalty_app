# Скрипт для фиксинга параметров документа ВКР
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

vkr_path = "vkr/build/VKR_Villa_Jaconda.docx"
output_path = "vkr/build/VKR_Villa_Jaconda_FIXED.docx"

print("=" * 80)
print("ФИКСИНГ ПАРАМЕТРОВ ДОКУМЕНТА ВКР")
print("=" * 80)

doc = Document(vkr_path)

# 1. Фиксинг полей страницы
print("\n[1] Фиксинг полей страницы...")
for section in doc.sections:
    print(f"  Было: верх={section.top_margin.cm:.2f}, низ={section.bottom_margin.cm:.2f}, "
          f"лево={section.left_margin.cm:.2f}, право={section.right_margin.cm:.2f} см")
    
    # Устанавливаем правильные поля
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)  # Было 1.0, теперь 1.5
    
    print(f"  Стало: верх={section.top_margin.cm:.2f}, низ={section.bottom_margin.cm:.2f}, "
          f"лево={section.left_margin.cm:.2f}, право={section.right_margin.cm:.2f} см")

# 2. Замена размера шрифта с 14pt на 12pt
print("\n[2] Изменение размера шрифта с 14pt на 12pt...")
changed_runs = 0
for para in doc.paragraphs:
    for run in para.runs:
        if run.font.size and run.font.size.pt == 14.0:
            run.font.size = Pt(12)
            changed_runs += 1

print(f"  Изменено: {changed_runs} элементов текста")

# 3. Замена стилей Normal (Web) на Normal
print("\n[3] Замена стилей 'Normal (Web)' на 'Normal'...")
changed_styles = 0
for para in doc.paragraphs:
    if para.style and para.style.name == "Normal (Web)":
        # Применяем стиль Normal
        para.style = "Normal"
        changed_styles += 1

print(f"  Изменено: {changed_styles} параграфов со стилем 'Normal (Web)' на 'Normal'")

# 4. Убедимся что шрифт правильный для заголовков
print("\n[4] Проверка шрифтов заголовков...")
heading_count = 0
for para in doc.paragraphs:
    if para.style and 'Heading' in para.style.name:
        heading_count += 1
        for run in para.runs:
            if run.font.size is None or run.font.size.pt < 12:
                run.font.size = Pt(12)

print(f"  Проверено: {heading_count} заголовков")

# 5. Установка выравнивания по ширине для основного текста
print("\n[5] Установка выравнивания по ширине...")
aligned_count = 0
for para in doc.paragraphs:
    if para.style and 'Normal' in para.style.name:
        if para.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            aligned_count += 1

print(f"  Выравнено: {aligned_count} параграфов")

# Сохраняем документ
print("\n[6] Сохранение документа...")
doc.save(output_path)
print(f"  Сохранено в: {output_path}")

print("\n" + "=" * 80)
print("РЕЗУЛЬТАТЫ ФИКСИНГА:")
print("=" * 80)
print(f"""
ИСПРАВЛЕНО:
1. Правое поле: 1.0 см -> 1.5 см
2. Размер шрифта: 14pt -> 12pt ({changed_runs} элементов)
3. Стили: 'Normal (Web)' -> 'Normal' ({changed_styles} параграфов)
4. Выравнивание: установлено по ширине для основного текста
5. Шрифты заголовков: проверены и выставлены минимум 12pt

ФАЙЛ СОХРАНЕН: vkr/build/VKR_Villa_Jaconda_FIXED.docx
""")
